# ReportGenerator GUI Tool
# Creato da Alessandro Frullo

import io
import threading
import tkinter as tk
from functools import lru_cache
from pathlib import Path
from tkinter import ttk, filedialog, messagebox, simpledialog

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENTATION
from PIL import Image, ImageChops, ImageDraw, ImageFont

__author__ = "Alessandro Frullo"

OPERATORS = ['Iliad', 'TIM', 'VF', 'W3']
ORDER = [
    'GSM900 RXLEV', 'LTE800 RSRP', 'LTE800 QUAL', 'UMTS900 RSCP', 'UMTS900 QUAL',
    'LTE1800 RSRP', 'LTE1800 QUAL', 'LTE2100 RSRP', 'LTE2100 QUAL', 'LTE2100 RSRP B100',
    'LTE RSRQ B100', 'UMTS2100 RSCP', 'UMTS2100 QUAL', 'LTE2600 RSRP', 'LTE2600 QUAL',
    'RSRQ 700', 'RSRP 700', 'RSRP 3500', 'RSRQ 3500',
    '5G SS-RSRP', '5G SS-RSRQ'
]
# Versioni minuscole pre-calcolate, usate per l'ordinamento.
_ORDER_LC = [lbl.lower() for lbl in ORDER]

# Color palette
MAIN_BG = '#f2f2f2'         # Light grey background
ACCENT_ORANGE = '#FFA500'   # Orange accent
ACCENT_BLUE = '#003366'     # Dark blue accent

# Dimensioni layout documento (in pollici)
SECTION_MARGIN = 0.2
HEADING_SPACE_AFTER = 0.2


# ---------------------------------------------------------------------------
# Utilità immagini
# ---------------------------------------------------------------------------

def extract_label_name(image_path):
    """Ricava l'etichetta 'OPERATORE TECNOLOGIA' dal nome del file."""
    filename = Path(image_path).stem
    for op in OPERATORS:
        if op in filename:
            tech = filename.replace('_Workbook_', '').replace(op, '').strip()
            return f"{op} {tech}"
    return filename


def _crop(img, mode):
    """Ritaglia i bordi bianchi dell'immagine.

    Usa il bounding-box nativo di Pillow (implementato in C) invece di
    iterare i pixel in Python: stesso risultato, molto più veloce.

    mode: 'sides' (solo sinistra/destra), 'topbottom' (solo alto/basso)
          o 'both' (tutti e quattro i lati).
    """
    diff = ImageChops.difference(img, Image.new('RGB', img.size, (255, 255, 255)))
    bbox = diff.getbbox()  # (left, top, right, bottom) con right/bottom esclusivi
    if bbox is None:
        return img  # immagine completamente bianca: niente da ritagliare
    left, top, right, bottom = bbox
    w, h = img.size
    if mode == 'sides':
        return img.crop((left, 0, right, h))
    if mode == 'topbottom':
        return img.crop((0, top, w, bottom))
    return img.crop((left, top, right, bottom))  # 'both'


@lru_cache(maxsize=None)
def _get_font(size=26):
    """Carica (una sola volta) il font usato per le etichette."""
    try:
        return ImageFont.truetype('arial.ttf', size)
    except IOError:
        return ImageFont.load_default()


def _add_label(img, label):
    """Aggiunge una banda superiore con l'etichetta racchiusa in un riquadro."""
    font = _get_font()
    tw, th = font.getbbox(label)[2:]
    pad_x, pad_y = 20, 10
    band_h = th + 2 * pad_y
    new_w = max(img.width, tw + 2 * pad_x)
    new_img = Image.new('RGB', (new_w, img.height + band_h), 'white')
    draw = ImageDraw.Draw(new_img)
    draw.rectangle([(0, 0), (new_w, band_h)], outline=ACCENT_BLUE, width=2)
    draw.text(((new_w - tw) // 2, pad_y), label, fill=ACCENT_BLUE, font=font)
    new_img.paste(img, (0, band_h))
    return new_img


def load_processed_image(path, crop_mode, add_label):
    """Restituisce (source, (width, height)) pronto per `doc.add_picture`.

    Se non serve alcuna modifica si passa direttamente il percorso originale
    (nessuna ri-codifica). Altrimenti l'immagine elaborata resta in memoria
    come stream PNG: niente file temporanei su disco.
    """
    if crop_mode == 'none' and not add_label:
        with Image.open(path) as im:
            return path, im.size

    with Image.open(path) as src:
        img = src.convert('RGB')
    if crop_mode != 'none':
        img = _crop(img, crop_mode)
    if add_label:
        img = _add_label(img, extract_label_name(path))

    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf, img.size


def sort_images_by_order(images):
    """Ordina le immagini secondo la sequenza definita in ORDER."""
    def key(n):
        stem = Path(n).stem.lower()
        for i, lbl in enumerate(_ORDER_LC):
            if lbl in stem:
                return i
        return len(_ORDER_LC)
    return sorted(images, key=key)


def add_images_to_doc(doc, title, imgs, crop_mode, add_label):
    """Aggiunge una sezione con titolo e tutte le immagini ordinate."""
    sec = doc.add_section()
    sec.top_margin = sec.bottom_margin = Inches(SECTION_MARGIN)
    sec.left_margin = sec.right_margin = Inches(SECTION_MARGIN)
    max_w = sec.page_width.inches - 2 * SECTION_MARGIN
    doc.add_heading(title, level=2).paragraph_format.space_after = Inches(HEADING_SPACE_AFTER)

    max_h = sec.page_height.inches - 2 * SECTION_MARGIN
    for path in sort_images_by_order(imgs):
        try:
            source, (w, h) = load_processed_image(path, crop_mode, add_label)
            if w / h > max_w / max_h:
                final_w, final_h = max_w, max_w * h / w
            else:
                final_w, final_h = max_h * w / h, max_h
            doc.add_picture(source, width=Inches(final_w), height=Inches(final_h))
        except Exception as e:
            print(f"Errore con {path}: {e}")


# ---------------------------------------------------------------------------
# Applicazione GUI
# ---------------------------------------------------------------------------

class DocGeneratorApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title('ReportGenerator - Selektra Italia')
        self.geometry('1000x810')
        self.configure(bg=MAIN_BG)

        # Styles
        style = ttk.Style(self)
        style.theme_use('clam')
        style.configure('TLabel', background=MAIN_BG, foreground=ACCENT_BLUE, font=('Arial', 11))
        style.configure('TButton', background=ACCENT_BLUE, foreground='white', font=('Arial', 11), padding=6)
        style.map('TButton', background=[('active', ACCENT_ORANGE)])
        style.configure('Primary.TButton', background=ACCENT_ORANGE, foreground=ACCENT_BLUE,
                        font=('Arial', 12, 'bold'), padding=10)
        style.map('Primary.TButton', background=[('active', '#e09400'), ('disabled', '#cccccc')])
        style.configure('TEntry', fieldbackground='white', padding=4)
        style.configure('TLabelframe', background=MAIN_BG, foreground=ACCENT_BLUE, font=('Arial', 11, 'bold'), padding=8)
        style.configure('TLabelframe.Label', background=MAIN_BG, foreground=ACCENT_BLUE, font=('Arial', 11, 'bold'))
        style.configure('Footer.TLabel', background=MAIN_BG, foreground='#888888', font=('Arial', 9, 'italic'))
        style.configure('Horizontal.TProgressbar', background=ACCENT_ORANGE, troughcolor='#e0e0e0')

        # Header banner
        header = tk.Frame(self, bg=ACCENT_BLUE, height=65)
        header.pack(fill='x')
        header.pack_propagate(False)
        tk.Label(header, text='ReportGenerator', bg=ACCENT_BLUE, fg='white',
                 font=('Arial', 22, 'bold')).pack(side='left', padx=20, pady=12)
        tk.Label(header, text='· Selektra Italia', bg=ACCENT_BLUE, fg=ACCENT_ORANGE,
                 font=('Arial', 13)).pack(side='left', pady=14)

        self.blocks = {}
        self._create_widgets()

    def _create_widgets(self):
        # Settings frame
        frame = ttk.Labelframe(self, text='Impostazioni')
        frame.pack(fill='x', padx=15, pady=10)
        frame.columnconfigure(1, weight=1)

        ttk.Label(frame, text='Titolo Documento:').grid(row=0, column=0, sticky='w', padx=5, pady=5)
        self.title_var = tk.StringVar()
        ttk.Entry(frame, textvariable=self.title_var).grid(row=0, column=1, sticky='ew', padx=5)

        ttk.Label(frame, text='Modalità Ritaglio:').grid(row=1, column=0, sticky='w', padx=5, pady=5)
        self.crop_var = tk.StringVar(value='none')
        ttk.Combobox(frame, textvariable=self.crop_var, state='readonly',
                     values=['none', 'sides', 'topbottom', 'both']).grid(row=1, column=1, sticky='w', padx=5)

        self.label_var = tk.BooleanVar()
        ttk.Checkbutton(frame, text='Aggiungi Etichetta', variable=self.label_var).grid(row=2, column=1, sticky='w', padx=5, pady=5)

        ttk.Label(frame, text='Cartella di Output:').grid(row=3, column=0, sticky='w', padx=5, pady=5)
        self.out_var = tk.StringVar()
        ttk.Entry(frame, textvariable=self.out_var).grid(row=3, column=1, sticky='ew', padx=5)
        ttk.Button(frame, text='Sfoglia...', command=self._select_output_folder).grid(row=3, column=2, padx=5)

        ttk.Label(frame, text='Sottocartelle:').grid(row=4, column=0, sticky='nw', padx=5, pady=5)
        sub_frame = ttk.Frame(frame)
        sub_frame.grid(row=4, column=1, columnspan=2, sticky='ew', padx=5, pady=5)
        self.subfolders_list = tk.Listbox(sub_frame, font=('Arial', 10), height=5,
                                          selectmode='browse', activestyle='none',
                                          fg=ACCENT_BLUE, bg='white', relief='flat',
                                          selectbackground=ACCENT_ORANGE, selectforeground='white',
                                          highlightbackground='#cccccc', highlightthickness=1)
        self.subfolders_list.pack(side='left', fill='both', expand=True)
        sub_scroll = ttk.Scrollbar(sub_frame, orient='vertical', command=self.subfolders_list.yview)
        sub_scroll.pack(side='right', fill='y')
        self.subfolders_list.config(yscrollcommand=sub_scroll.set)

        # Image blocks frame
        blk_frame = ttk.Labelframe(self, text='Blocchi di Immagini')
        blk_frame.pack(fill='both', expand=True, padx=15, pady=10)
        blk_frame.rowconfigure(0, weight=1)
        blk_frame.columnconfigure(0, weight=1)

        list_container = ttk.Frame(blk_frame)
        list_container.grid(row=0, column=0, sticky='nsew', padx=5, pady=5)
        self.blocks_list = tk.Listbox(list_container, font=('Arial', 10),
                                      selectbackground=ACCENT_ORANGE, selectforeground='white',
                                      activestyle='none', relief='flat',
                                      highlightbackground='#cccccc', highlightthickness=1)
        self.blocks_list.pack(side='left', fill='both', expand=True)
        scrollbar = ttk.Scrollbar(list_container, orient='vertical', command=self.blocks_list.yview)
        scrollbar.pack(side='right', fill='y')
        self.blocks_list.config(yscrollcommand=scrollbar.set)

        btn_frame = ttk.Frame(blk_frame)
        btn_frame.grid(row=0, column=1, sticky='ns', padx=5)
        ttk.Button(btn_frame, text='Aggiungi Blocco', command=self._add_block).pack(fill='x', pady=5)
        ttk.Button(btn_frame, text='Rimuovi Blocco', command=self._remove_block).pack(fill='x', pady=5)

        # Generate and progress frame
        gen_frame = ttk.Frame(self)
        gen_frame.pack(fill='x', padx=15, pady=10)
        self.generate_btn = ttk.Button(gen_frame, text='Genera Documenti',
                                       command=self._generate_documents, style='Primary.TButton')
        self.generate_btn.pack(side='left')
        self.progress = ttk.Progressbar(gen_frame, style='Horizontal.TProgressbar', orient='horizontal',
                                         mode='determinate', maximum=100)
        self.progress.pack(side='left', fill='x', expand=True, padx=10)

        # Footer label
        footer = ttk.Label(self, text='Creato da Alessandro Frullo', style='Footer.TLabel')
        footer.pack(side='bottom', pady=5)

    def _select_output_folder(self):
        folder = filedialog.askdirectory(title='Seleziona Cartella di Output')
        if folder:
            self.out_var.set(folder)
            self._refresh_subfolders(folder)

    def _refresh_subfolders(self, folder):
        self.subfolders_list.delete(0, 'end')
        subfolders = sorted(p for p in Path(folder).iterdir() if p.is_dir())
        if not subfolders:
            self.subfolders_list.insert('end', '(nessuna sottocartella trovata)')
            return
        for path in subfolders:
            try:
                child_names = {s.name.lower() for s in path.iterdir() if s.is_dir()}
            except PermissionError:
                child_names = set()
            has_4g = any('4g' in s for s in child_names)
            has_5g = any('5g' in s for s in child_names)
            self.subfolders_list.insert('end', path.name)
            idx = self.subfolders_list.size() - 1
            if has_4g and has_5g:
                self.subfolders_list.itemconfig(idx, foreground='white', background=ACCENT_ORANGE,
                                                selectbackground='#e09400', selectforeground='white')
            elif has_5g:
                self.subfolders_list.itemconfig(idx, foreground='white', background=ACCENT_BLUE,
                                                selectbackground='#002244', selectforeground='white')

    def _add_block(self):
        title = simpledialog.askstring('Titolo Blocco', 'Inserisci il titolo del blocco:')
        if not title:
            return
        files = filedialog.askopenfilenames(
            title=f'Seleziona immagini per: {title}',
            filetypes=[('Images', '*.jpg;*.png;*.jpeg;*.bmp;*.gif')],
        )
        if files:
            self.blocks[title] = list(files)
            self.blocks_list.insert('end', title)

    def _remove_block(self):
        sel = self.blocks_list.curselection()
        if not sel:
            return
        title = self.blocks_list.get(sel)
        del self.blocks[title]
        self.blocks_list.delete(sel)

    def _generate_documents(self):
        title = self.title_var.get().strip().replace(' ', '_')
        out_dir = Path(self.out_var.get() or Path.cwd())
        crop_mode = self.crop_var.get()
        add_label = self.label_var.get()
        if not title:
            messagebox.showerror('Errore', 'Titolo obbligatorio')
            return
        if not self.blocks:
            messagebox.showerror('Errore', 'Aggiungi almeno un blocco')
            return

        # Disable UI and reset progress
        self.generate_btn.config(state='disabled')
        self.progress['value'] = 0

        # Run generation in background thread
        threading.Thread(
            target=self._run_generation,
            args=(title, out_dir, crop_mode, add_label),
            daemon=True,
        ).start()

    def _run_generation(self, title, out_dir, crop_mode, add_label):
        try:
            out_dir.mkdir(parents=True, exist_ok=True)
            ops = {op: out_dir / f'{title}_{op}.docx' for op in OPERATORS}
            # Stem in minuscolo pre-calcolato una sola volta per ogni immagine.
            stems = {p: Path(p).stem.lower() for pics in self.blocks.values() for p in pics}
            total = len(OPERATORS)

            for i, op in enumerate(OPERATORS):
                op_lc = op.lower()
                # Per ogni blocco, le immagini che appartengono a questo operatore.
                blocks_for_op = [
                    (blk, [p for p in pics if op_lc in stems[p]])
                    for blk, pics in self.blocks.items()
                ]
                blocks_for_op = [(blk, rel) for blk, rel in blocks_for_op if rel]
                if not blocks_for_op:
                    self.after(0, lambda v=int((i + 1) / total * 100): self._set_progress(v))
                    continue

                path = ops[op]
                exists = path.exists()
                doc = Document(str(path)) if exists else Document()
                if not exists:
                    style = doc.styles['Normal']
                    style.font.name = 'Arial'
                    style.font.size = Pt(12)
                    sec0 = doc.sections[0]
                    sec0.orientation = WD_ORIENTATION.LANDSCAPE
                    sec0.page_width, sec0.page_height = sec0.page_height, sec0.page_width

                for blk, rel in blocks_for_op:
                    add_images_to_doc(doc, blk, rel, crop_mode, add_label)
                doc.save(str(path))
                self.after(0, lambda v=int((i + 1) / total * 100): self._set_progress(v))

            self.after(0, lambda: self._on_complete(out_dir))
        except Exception as e:
            self.after(0, lambda err=e: self._on_error(err))

    def _set_progress(self, value):
        self.progress['value'] = value

    def _on_complete(self, out_dir):
        self.progress['value'] = 100
        self.generate_btn.config(state='normal')
        messagebox.showinfo('Successo', f'Documenti creati in: {out_dir}')

    def _on_error(self, error):
        self.progress['value'] = 0
        self.generate_btn.config(state='normal')
        messagebox.showerror('Errore', f'Generazione non riuscita: {error}')


if __name__ == '__main__':
    app = DocGeneratorApp()
    app.mainloop()
