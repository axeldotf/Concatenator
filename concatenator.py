# Copyright 2025 Alessandro Frullo
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

import os
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENTATION
from PIL import Image, ImageDraw, ImageFont
from tqdm import tqdm
import threading

__author__ = "Alessandro Frullo"

OPERATORS = ['Iliad', 'TIM', 'VF', 'W3']
ORDER = [
    'GSM900 RXLEV', 'LTE800 RSRP', 'LTE800 QUAL', 'UMTS900 RSCP', 'UMTS900 QUAL',
    'LTE1800 RSRP', 'LTE1800 QUAL', 'LTE2100 RSRP', 'LTE2100 QUAL', 'LTE2100 RSRP B100',
    'LTE RSRQ B100', 'UMTS2100 RSCP', 'UMTS2100 QUAL', 'LTE2600 RSRP', 'LTE2600 QUAL',
    'RSRQ 700', 'RSRP 700', 'RSRP 3500', 'RSRQ 3500'
]

# Color palette
MAIN_BG = '#f2f2f2'         # Light grey background
ACCENT_ORANGE = '#FFA500'   # Orange accent
ACCENT_BLUE = '#003366'     # Dark blue accent

# Label extraction utility
def extract_label_name(image_path):
    filename = Path(image_path).stem
    for op in OPERATORS:
        if op in filename:
            tech = filename.replace('_Workbook_', '').replace(op, '').strip()
            return f"{op} {tech}"
    return filename

# Cropping helpers
def crop_sides(img):
    pixels = img.load()
    w, h = img.size
    left = next(x for x in range(w) if any(pixels[x, y] != (255, 255, 255) for y in range(h)))
    right = next(x for x in range(w-1, -1, -1) if any(pixels[x, y] != (255, 255, 255) for y in range(h))) + 1
    return img.crop((left, 0, right, h))

def crop_top_bottom(img):
    pixels = img.load()
    w, h = img.size
    top = next(y for y in range(h) if any(pixels[x, y] != (255, 255, 255) for x in range(w)))
    bottom = next(y for y in range(h-1, -1, -1) if any(pixels[x, y] != (255, 255, 255) for x in range(w))) + 1
    return img.crop((0, top, w, bottom))

# Process image based on user options
def process_image(path, crop_mode, add_label):
    with Image.open(path) as img:
        img = img.convert('RGB')
        if crop_mode in ('sides', 'both'):
            img = crop_sides(img)
        if crop_mode in ('topbottom', 'both'):
            img = crop_top_bottom(img)
        if add_label:
            label = extract_label_name(path)
            try:
                font = ImageFont.truetype('arial.ttf', 26)
            except IOError:
                font = ImageFont.load_default()
            tw, th = font.getbbox(label)[2:]
            pad_x, pad_y = 20, 10
            new_w = max(img.width, tw + 2 * pad_x)
            new_h = img.height + th + 2 * pad_y
            new_img = Image.new('RGB', (new_w, new_h), 'white')
            draw = ImageDraw.Draw(new_img)
            draw.rectangle([(0, 0), (new_w, th + 2 * pad_y)], outline=ACCENT_BLUE, width=2)
            draw.text(((new_w - tw) // 2, pad_y), label, fill=ACCENT_BLUE, font=font)
            new_img.paste(img, (0, th + 2 * pad_y))
            img = new_img
        output = path
        if crop_mode != 'none' or add_label:
            suffix_parts = []
            if crop_mode != 'none':
                suffix_parts.append(crop_mode)
            if add_label:
                suffix_parts.append('labeled')
            output = path.replace('.', '_' + '_'.join(suffix_parts) + '.', 1)
            img.save(output)
        return output

# Sort images by predefined order
def sort_images_by_order(images):
    def key(n):
        stem = Path(n).stem.lower()
        for i, lbl in enumerate(ORDER):
            if lbl.lower() in stem:
                return i
        return len(ORDER)
    return sorted(images, key=key)

# Add images into Word document
def add_images_to_doc(doc, title, imgs, crop_mode, add_label):
    sec = doc.add_section()
    sec.top_margin = sec.bottom_margin = Inches(0.2)
    sec.left_margin = sec.right_margin = Inches(0.2)
    maxw = sec.page_width.inches - 0.4
    doc.add_heading(title, level=2).paragraph_format.space_after = Inches(0.2)
    for img in sort_images_by_order(imgs):
        try:
            proc = process_image(img, crop_mode, add_label)
            with Image.open(proc) as im:
                w, h = im.size
                doc.add_picture(proc, width=Inches(maxw), height=Inches(maxw * (h / w)))
            if proc != img:
                os.remove(proc)
        except Exception as e:
            print(f"Errore con {img}: {e}")

# Main GUI Application
class DocGeneratorApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title('ReportGenerator - Selektra Italia')
        self.geometry('1000x750')
        self.configure(bg=MAIN_BG)

        # Styles
        style = ttk.Style(self)
        style.theme_use('clam')
        style.configure('TLabel', background=MAIN_BG, foreground=ACCENT_BLUE, font=('Arial', 11))
        style.configure('TButton', background=ACCENT_BLUE, foreground='white', font=('Arial', 11), padding=6)
        style.map('TButton', background=[('active', ACCENT_ORANGE)])
        style.configure('TEntry', fieldbackground='white', padding=4)
        style.configure('TLabelframe', background=MAIN_BG, foreground=ACCENT_BLUE, font=('Arial', 12, 'bold'), padding=6)
        style.configure('Footer.TLabel', background=MAIN_BG, foreground=ACCENT_BLUE, font=('Arial', 9, 'italic'))
        style.configure('Horizontal.TProgressbar', background=ACCENT_ORANGE)

        self.blocks = {}
        self._create_widgets()

    def _create_widgets(self):
        # Settings frame
        frame = ttk.Labelframe(self, text='Impostazioni')
        frame.pack(fill='x', padx=15, pady=10)
        frame.columnconfigure(1, weight=1)

        ttk.Label(frame, text='Titolo Documento:').grid(row=0, column=0, sticky='w', padx=5, pady=5)
        self.title_var = tk.StringVar()
        ttk.Entry(frame, textvariable=self.title_var).grid(row=0, column=1, sticky='ew', padx=5)

        ttk.Label(frame, text='Modalità Ritaglio:').grid(row=1, column=0, sticky='w', padx=5, pady=5)
        self.crop_var = tk.StringVar(value='none')
        ttk.Combobox(frame, textvariable=self.crop_var, state='readonly', values=['none', 'sides', 'topbottom', 'both']).grid(row=1, column=1, sticky='w', padx=5)

        self.label_var = tk.BooleanVar()
        ttk.Checkbutton(frame, text='Aggiungi Etichetta', variable=self.label_var).grid(row=2, column=1, sticky='w', padx=5, pady=5)

        ttk.Label(frame, text='Cartella di Output:').grid(row=3, column=0, sticky='w', padx=5, pady=5)
        self.out_var = tk.StringVar()
        ttk.Entry(frame, textvariable=self.out_var).grid(row=3, column=1, sticky='ew', padx=5)
        ttk.Button(frame, text='Sfoglia...', command=self._select_output_folder).grid(row=3, column=2, padx=5)

        # Image blocks frame
        blk_frame = ttk.Labelframe(self, text='Blocchi di Immagini')
        blk_frame.pack(fill='both', expand=True, padx=15, pady=10)
        blk_frame.rowconfigure(0, weight=1)
        blk_frame.columnconfigure(0, weight=1)

        list_container = ttk.Frame(blk_frame)
        list_container.grid(row=0, column=0, sticky='nsew', padx=5, pady=5)
        self.blocks_list = tk.Listbox(list_container, font=('Arial', 10))
        self.blocks_list.pack(side='left', fill='both', expand=True)
        scrollbar = ttk.Scrollbar(list_container, orient='vertical', command=self.blocks_list.yview)
        scrollbar.pack(side='right', fill='y')
        self.blocks_list.config(yscrollcommand=scrollbar.set)

        btn_frame = ttk.Frame(blk_frame)
        btn_frame.grid(row=0, column=1, sticky='ns', padx=5)
        ttk.Button(btn_frame, text='Aggiungi Blocco', command=self._add_block).pack(fill='x', pady=5)
        ttk.Button(btn_frame, text='Rimuovi Blocco', command=self._remove_block).pack(fill='x', pady=5)

        # Generate and progress frame
        gen_frame = ttk.Frame(self)
        gen_frame.pack(fill='x', padx=15, pady=10)
        self.generate_btn = ttk.Button(gen_frame, text='Genera Documenti', command=self._generate_documents)
        self.generate_btn.pack(side='left')
        self.progress = ttk.Progressbar(gen_frame, style='Horizontal.TProgressbar', orient='horizontal', mode='indeterminate')
        self.progress.pack(side='left', fill='x', expand=True, padx=10)

        # Footer label
        footer = ttk.Label(self, text='Creato da Alessandro Frullo', style='Footer.TLabel')
        footer.pack(side='bottom', pady=5)

    def _select_output_folder(self):
        folder = filedialog.askdirectory(title='Seleziona Cartella di Output')
        if folder:
            self.out_var.set(folder)

    def _add_block(self):
        title = simpledialog.askstring('Titolo Blocco', 'Inserisci il titolo del blocco:')
        if not title:
            return
        files = filedialog.askopenfilenames(title=f'Seleziona immagini per: {title}', filetypes=[('Images', '*.jpg;*.png;*.jpeg;*.bmp;*.gif')])
        if files:
            self.blocks[title] = list(files)
            self.blocks_list.insert('end', title)

    def _remove_block(self):
        sel = self.blocks_list.curselection()
        if not sel:
            return
        title = self.blocks_list.get(sel)
        del self.blocks[title]
        self.blocks_list.delete(sel)

    def _generate_documents(self):
        title = self.title_var.get().strip().replace(' ', '_')
        out_dir = Path(self.out_var.get() or Path.cwd())
        crop_mode = self.crop_var.get()
        add_label = self.label_var.get()
        if not title:
            messagebox.showerror('Errore', 'Titolo obbligatorio')
            return
        if not self.blocks:
            messagebox.showerror('Errore', 'Aggiungi almeno un blocco')
            return

        # Disable UI and start progress
        self.generate_btn.config(state='disabled')
        self.progress.start(10)

        # Run generation in background thread
        threading.Thread(target=self._run_generation, args=(title, out_dir, crop_mode, add_label), daemon=True).start()

    def _run_generation(self, title, out_dir, crop_mode, add_label):
        out_dir.mkdir(parents=True, exist_ok=True)
        ops = {op: out_dir / f'{title}_{op}.docx' for op in OPERATORS}
        images = [img for imgs in self.blocks.values() for img in imgs]

        for op in OPERATORS:
            imgs = [i for i in images if op.lower() in Path(i).stem.lower()]
            if not imgs:
                continue
            doc = Document(ops[op]) if Path(ops[op]).exists() else Document()
            if not Path(ops[op]).exists():
                style = doc.styles['Normal']
                style.font.name = 'Arial'
                style.font.size = Pt(12)
                sec0 = doc.sections[0]
                sec0.orientation = WD_ORIENTATION.LANDSCAPE
                sec0.page_width, sec0.page_height = sec0.page_height, sec0.page_width
            for blk, pics in self.blocks.items():
                rel = [i for i in pics if i in imgs]
                if rel:
                    add_images_to_doc(doc, blk, rel, crop_mode, add_label)
            doc.save(ops[op])

        # Upon completion, update UI in main thread
        def on_complete():
            self.progress.stop()
            self.generate_btn.config(state='normal')
            messagebox.showinfo('Successo', f'Documenti creati in: {out_dir}')
        self.after(0, on_complete)

if __name__ == '__main__':
    app = DocGeneratorApp()
    app.mainloop()
